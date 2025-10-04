import re
import uuid
from pathlib import Path
from docx import Document
from django.conf import settings

# Находим ЛЮБОЕ содержимое в { ... }, кроме вложенных фигурных скобок
TOKEN_RE = re.compile(r"\{([^{}]+)\}")

def _norm(s: str) -> str:
    """Нормализация ключа: убрать все пробелы и привести к нижнему регистру."""
    return re.sub(r"\s+", "", str(s)).lower()

def _collect_tokens_from_paragraphs(paragraphs) -> set[str]:
    tokens = set()
    for p in paragraphs:
        for m in TOKEN_RE.finditer(p.text or ""):
            tokens.add(m.group(0))  # целиком: '{...}'
    return tokens

def _collect_tokens_from_tables(tables) -> set[str]:
    tokens = set()
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                tokens |= _collect_tokens_from_paragraphs(cell.paragraphs)
    return tokens

def _replace_in_paragraphs(paragraphs, mapping_exact: dict[str, str]):
    for p in paragraphs:
        text = p.text or ""
        if not text:
            continue
        for token, value in mapping_exact.items():
            if token in text:
                text = text.replace(token, value)
        p.text = text  # да, это пересоздаёт runs — зато надёжно заменяет даже «разбитые» плейсхолдеры

def _replace_in_tables(tables, mapping_exact: dict[str, str]):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, mapping_exact)

def generate_document(template_path: str, values_dict: dict[str, str], default_placeholder: str = "—") -> Path:
    """
    Открывает .docx и заменяет { ... } на значения из values_dict.
    — Поддерживает кириллицу, пробелы и произвольный регистр внутри { }.
    — Если значение не найдено, подставляет default_placeholder ('—').
    — Заменяет в тексте, таблицах, header/footer.
    Возвращает путь к сохранённому файлу в MEDIA_ROOT/generated/<uuid>.docx.
    """
    doc = Document(template_path)

    # 1) Собираем все ТЕКСТОВЫЕ ТОКЕНЫ вида '{...}' из документа
    tokens: set[str] = set()
    tokens |= _collect_tokens_from_paragraphs(doc.paragraphs)
    tokens |= _collect_tokens_from_tables(doc.tables)

    # Также пройдёмся по шапкам и подвалам
    for section in doc.sections:
        hdr, ftr = section.header, section.footer
        tokens |= _collect_tokens_from_paragraphs(hdr.paragraphs)
        tokens |= _collect_tokens_from_tables(hdr.tables)
        tokens |= _collect_tokens_from_paragraphs(ftr.paragraphs)
        tokens |= _collect_tokens_from_tables(ftr.tables)

    # 2) Нормализуем словарь значений для кейс-/пробел-инвариантного поиска
    values_norm = { _norm(k): str(v) for k, v in (values_dict or {}).items() }

    # 3) Готовим точные замены: '{оригинал_из_дока}' -> 'значение'
    mapping_exact: dict[str, str] = {}
    for token in tokens:
        inner = token[1:-1]  # содержимое без скобок
        key_norm = _norm(inner)
        value = values_norm.get(key_norm, default_placeholder)
        mapping_exact[token] = value

    # 4) Применяем замены
    _replace_in_paragraphs(doc.paragraphs, mapping_exact)
    _replace_in_tables(doc.tables, mapping_exact)

    for section in doc.sections:
        hdr, ftr = section.header, section.footer
        _replace_in_paragraphs(hdr.paragraphs, mapping_exact)
        _replace_in_tables(hdr.tables, mapping_exact)
        _replace_in_paragraphs(ftr.paragraphs, mapping_exact)
        _replace_in_tables(ftr.tables, mapping_exact)

    # 5) Сохраняем
    out_dir = Path(settings.MEDIA_ROOT) / "generated"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{uuid.uuid4().hex}.docx"
    doc.save(out_path)
    return out_path
